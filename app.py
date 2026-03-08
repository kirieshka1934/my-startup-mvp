import streamlit as st
import requests
import base64
from PIL import Image
from docx import Document
import io

# --- НАСТРОЙКИ ---
st.set_page_config(page_title="StudyFlow AI", layout="wide")
st.title("🎓 StudyFlow: Фото → Конспект")

# --- 1. ФУНКЦИЯ РАСПОЗНАВАНИЯ ТЕКСТА (OCR) ---
def recognize_text(image_bytes, api_key, folder_id):
    encoded_image = base64.b64encode(image_bytes).decode('utf-8')
    url = "https://ocr.api.cloud.yandex.net/ocr/v1/recognizeText"
    headers = {"Content-Type": "application/json", "Authorization": f"Api-Key {api_key}"}
    payload = {
        "mimeType": "image/jpeg",
        "languageCodes": ["*"],
        "model": "page",
        "content": encoded_image
    }
    response = requests.post(url, headers=headers, json=payload)
    if response.status_code == 200:
        # Собираем все кусочки текста в одну строку
        blocks = response.json()['result']['textAnnotation']['blocks']
        full_text = ""
        for block in blocks:
            for line in block['lines']:
                full_text += line['text'] + " "
        return full_text
    else:
        return None

# --- 2. ФУНКЦИЯ ОБРАБОТКИ ТЕКСТА (GPT) ---
def ask_yandex_gpt(text, api_key, folder_id):
    url = "https://llm.api.cloud.yandex.net/foundationModels/v1/completion"
    headers = {"Content-Type": "application/json", "Authorization": f"Api-Key {api_key}"}
    payload = {
        "modelUri": f"gpt://{folder_id}/yandexgpt-lite",
        "completionOptions": {"stream": False, "temperature": 0.5, "maxTokens": "2000"},
        "messages": [
            {"role": "system", "text": "Ты — AI-тьютор. Сделай структурированный конспект и 3 вопроса по тексту."},
            {"role": "user", "text": text}
        ]
    }
    response = requests.post(url, headers=headers, json=payload)
    if response.status_code == 200:
        return response.json()['result']['alternatives'][0]['message']['text']
    return f"Ошибка GPT: {response.text}"

# --- 3. ФУНКЦИЯ WORD ---
def create_docx(text):
    doc = Document()
    doc.add_heading('Конспект StudyFlow', 0)
    for line in text.split('\n'):
        doc.add_paragraph(line)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- ИНТЕРФЕЙС ---
uploaded_file = st.file_uploader("Загрузи фото конспекта (JPG, PNG):", type=["jpg", "jpeg", "png"])
user_text = st.text_area("Или просто вставь текст вручную:", height=150)

if st.button("🚀 Обработать материал"):
    api_key = st.secrets["YC_API_KEY"]
    folder_id = st.secrets["YC_FOLDER_ID"]
    
    source_text = ""
    
    if uploaded_file:
        with st.spinner("Распознаю текст с фото..."):
            img_bytes = uploaded_file.read()
            st.image(img_bytes, width=400)
            source_text = recognize_text(img_bytes, api_key, folder_id)
            if not source_text:
                st.error("Не удалось прочитать текст на фото.")
    elif user_text:
        source_text = user_text
        
    if source_text:
        with st.spinner("Нейросеть пишет конспект..."):
            final_result = ask_yandex_gpt(source_text, api_key, folder_id)
            st.markdown("### Ваш результат:")
            st.markdown(final_result)
            
            file = create_docx(final_result)
            st.download_button("📥 Скачать Word", file, "konspekt.docx")
    else:
        st.warning("Загрузи фото или вставь текст!")
